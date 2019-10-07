# OGARNIJ TWORZENIE NOWEJ NOTYFIKACJI I WYSYLANIE MAILA (OUTPUT DOBRY W DOCX STARY W OUTLOOKU, POPRAWIC)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from os import environ, getcwd
import docx, pyperclip, re, os, getpass, sys, subprocess


userName = getpass.getuser() 

# Lists all *.msg files located in MSG directory, returns chosen file name. 
# If there is a duplicate incident f.e: "P2_INC222" and "PRIORITY 2 - SNOW REF INC222" first one will be replaced with the second one
# Replaces the white spaces in file name with "_" because later VBS has issues with file permissions (don't know why though)
def choose_file():
    
    def checkForSpaces(x):
        if " " in x:
            return x.replace(" ", "_")
        else:
            return x
        
    def renameFile(old_name, new_name):
        try:
            os.rename(os.path.join(os.getcwd(), "Messages", file), os.path.join(os.getcwd(), "Messages", new_name))
            arr.append(new_name)
        except FileExistsError:
            os.remove(os.path.join(os.getcwd(), "Messages", new_name))
            os.rename(os.path.join(os.getcwd(), "Messages", file), os.path.join(os.getcwd(), "Messages", new_name))
            
    text = "\nRead notification template from:\n---------------"
    arr = ["output.docx"]
    
    for file in os.listdir(os.path.join(os.getcwd(), "Messages")):
        
        if file.endswith(".msg"):
            
            if re.search("^PRIORITY.+SNOW REF INC[0-9]{7}.+NOTIFICATION", file):
                new_name = "P" + [s for s in file.split() if s.isdigit()][0] + "_" + file[file.find('INC'):file.find('INC')+10] + ".msg"
                renameFile(file, new_name)
            
            else:
                new_name = checkForSpaces(file)
                renameFile(file, new_name)

    print(text)
    
    for i, name in enumerate(arr):
        print(str(i+1) + ". " + name)
    
    print("---------------")            
    fnameNo = input()
    
    if int(fnameNo) < 1 or int(fnameNo) > len(arr):
        print('\n!!!---You must input a number between 1 and ' + str(len(arr)) + '---!!!\n')
        choose_file()
    else:
        return arr[int(fnameNo) - 1]

# Returns Name and Surname
def getUserName():
    x = userName.split('.')
    name = x[0].capitalize()
    surname = x[1].capitalize()
    return (name + ' ' + surname)

def vbs_out(command):
    proc = subprocess.Popen(command, stdout = subprocess.DEVNULL, stderr = subprocess.PIPE, shell = True)
    print('\n---Wait for template to save in drafts---\n')
    proc.wait()
    if proc.returncode == 1:
        print('\n!!!!!!!!!!!!!!!!!!!\nAn error occured while trying to save the template in Outlook.\nClose Outlook and try again, in case this won\'t work, contact John\n!!!!!!!!!!!!!!!!!!!\n')
    else:
        print('\n---Template saved in Outlook -> CIM mailbox -> drafts folder---\n')
        
def vbs_in(command):
    proc = subprocess.Popen(command, stdout = subprocess.DEVNULL, stderr = subprocess.PIPE, shell = True)
    proc.wait()
    print('\n---Exporting data from mail---\n')
    while proc.poll() is None:
        pass
    if proc.returncode == 1:
        print('\n!!!!!!!!!!!!!!!!!!!\nAn error occured while trying to read the mail.\nClose Outlook and try again, in case this won\'t work, contact John\n!!!!!!!!!!!!!!!!!!!\n')
    else:
        print('\n---Export Complete---\n')
    

# Takes care of saving files
def save_file(doc, prio, incNo, stat, fname = 'output.msg'):
    
#     for file in os.listdir():
#         if file.endswith(".msg"):
#             os.rename(file, "output.msg")
    
    finalTouch(doc.tables[0])
    
    #GDZIES TUTAJ SIE SYPIE U KUBY
    try:
        if fname == "output.docx":
            doc.save('output.docx')
            print('\n---Template saved in output.docx---\n')
        else:
            doc.save('output.docx')
            print('\n---Template saved in output.docx---\n')
            vbs_out(os.path.join("template", "vbs", "out.vbs" + " " + fname + " " + prio + " " + incNo + " " + stat))
            
    
    except PermissionError:
        print('\n!!!---File in use, close output.docx and press ENTER to continue or type "stop" to cancel---!!!\n')
        x = input()
        if x == 'stop':
            pass
        else:
            save_file(doc, prio, incNo, stat, fname)
        
        
# Makes sure that the text is correctly formatted 
def finalTouch(tab):
    try:
        tab.cell(1,0).text = tab.cell(1,0).text + "\n"
        tab.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
        tab.rows[1].cells[0].paragraphs[0].runs[0].font.underline = True
        tab.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass

    for row in tab.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Calibri'
                    font.size = Pt(12)

###############################################################################################################################################
def paster():
    
# Searches for value given in the argument in the description. When found returns phrase until \n. When not founds returns -1
    def impacts():
        serviceImpact = input('''
Select Service Impact:
---------------
1. Messaging
2. Application
3. Network
4. Server
5. Workstation
6. Printer
7. Industrial Device
---------------        
''')

        if serviceImpact == '1':
            serviceImpact = 'Messaging'
            ciImpact = 'Messaging'            
        elif serviceImpact == '2':
            serviceImpact = 'Application'
            ciImpact = 'Application'
        elif serviceImpact == '3':
            serviceImpact = 'Network'
            ciImpact = 'Infrastructure'
        elif serviceImpact == '4':
            serviceImpact = 'Server'
            ciImpact = 'Infrastructure'
        elif serviceImpact == '5':
            serviceImpact = 'Workstation'
            ciImpact = 'Infrastructure'
        elif serviceImpact == '6':
            serviceImpact = 'Printer'
            ciImpact = 'Infrastructure'
        elif serviceImpact == '7':
            serviceImpact = 'Industrial Device'
            ciImpact = 'Industrial Device'
        else:
            print('\n!!!---You must input a number between 1 and 7---!!!\n')
            
        businessImpact = input('''
---------------
Select Business Impact
1. Production                   (Production line stopped)
2. Financial                    (Financial loss)
3. Shipping                     (train or truck blocked)
4. Safety                       (Employee safety)
5. Security                     (data or access security)
6. Multiple users from different locations are no able perform daily work
---------------
''')
    
        if businessImpact == '1':
            businessImpact = 'Production'
        elif businessImpact == '2':
            businessImpact = 'Financial'
        elif businessImpact == '3':
            businessImpact = 'Shipping'
        elif businessImpact == '4':
            businessImpact = 'Safety'
        elif businessImpact == '5':
            businessImpact = 'Security'
        elif businessImpact == '6':
            businessImpact = 'Multiple users from different locations are not able to perform daily work'
        else:
            print('\n!!!---You must input a number between 1 and 6---!!!\n')
            
        return [serviceImpact, ciImpact, businessImpact]
    
# Takes string x as argument, searches for x in INC detailed description and returns string that begins at x and ends at \n
    def parser(x):
        n = description.find(x)
        if x == 'ISSUE DESCRIPTION:':

            if description.find('[ENG]') != -1:
                n = description.find('[ENG]')
                return (description[n+5:] if n != -1 else -1)
            else:
                m = n+30            
                return (description[n:n+description[m:].find('\n')].split(':', 1)[1] if n != -1 else -1)
        
        else:
            return (description[n:n+description[n:].find('\n')].split(':', 1)[1] if n != -1 else -1)
        
# Assigning data from clipboard to variables, data must be a string containing "/nextEl," keyword 9 times in order for this function to work, if else it throws an error
    chFile = choose_file()
    data = pyperclip.paste()
    data = data.split('/nextEl,')
    
    if len(data) == 10:
        doc = docx.Document(os.path.join('template', 'template.docx'))
        table = doc.tables[0]
        incNo = data[0] 
        incStatus = data[1] 
        incPrio = data[2] 
        summary = data[3]
        description = data[4]
        RG = data[5]
        startDate = data[6]
        latestDate = data[7]
        latestUpdate = data[8]
        desc = parser('ISSUE DESCRIPTION:')
        location = parser('LOCATION')
        impact = impacts()
        
# Filling the table with scrapped values. If there's another table in *msg file print error
        try:
            table.cell(1,0).text = '\n' + 'P' + incPrio + ' [' + incNo + '] Incident Initial Notification' + '\n'
            table.cell(2,1).text = incNo
            table.cell(3,1).text = startDate 
            table.cell(4,1).text = 'P' + incPrio
            table.cell(4,3).text = incStatus
            table.cell(5,1).text = summary
            table.cell(6,1).text = (desc.strip() if desc != -1 else '')
            table.cell(7,1).text = impact[0]
            table.cell(7,3).text = impact[1]
            table.cell(8,1).text = impact[2]
            table.cell(10,1).text = (location if location != -1 else '')
            table.cell(11,1).text = getUserName()
            table.cell(11,3).text = RG
            table.cell(13,1).text = latestDate + ' - ' + latestUpdate
            table.cell(14,1).text = ('30 minutes' if incPrio == '1' else 'Upon Resolution')
        except IndexError:
            print('\n!!!---Make sure that in the file you are choosing is ONLY notification table---!!!\n')
         
        save_file(doc, incPrio, incNo, "INITIAL", chFile)
        
    else:
        print('\n!!!---Invalid data format, press ALT+5 in SNow and try again---!!!\n')
    
def colors(x):
    
# Fills the template with choosen colors, can also add latest work-notes update
    def filling(val):
        fill1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill3 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill4 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill5 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill6 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill7 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill8 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill9 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill10 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill11 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill12 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill13 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill14 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill15 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
        fill16 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), val))
            
        table.cell(2,0)._tc.get_or_add_tcPr().append(fill1)
        table.cell(3,0)._tc.get_or_add_tcPr().append(fill2)
        table.cell(4,0)._tc.get_or_add_tcPr().append(fill3)
        table.cell(4,2)._tc.get_or_add_tcPr().append(fill4)
        table.cell(5,0)._tc.get_or_add_tcPr().append(fill5)
        table.cell(6,0)._tc.get_or_add_tcPr().append(fill6)
        table.cell(7,0)._tc.get_or_add_tcPr().append(fill7)
        table.cell(7,2)._tc.get_or_add_tcPr().append(fill8)
        table.cell(8,0)._tc.get_or_add_tcPr().append(fill9)
        table.cell(9,0)._tc.get_or_add_tcPr().append(fill10)
        table.cell(10,0)._tc.get_or_add_tcPr().append(fill11)
        table.cell(11,0)._tc.get_or_add_tcPr().append(fill12)
        table.cell(11,2)._tc.get_or_add_tcPr().append(fill13)
        table.cell(12,0)._tc.get_or_add_tcPr().append(fill14)
        table.cell(13,0)._tc.get_or_add_tcPr().append(fill15)        
        table.cell(14,0)._tc.get_or_add_tcPr().append(fill16)
    
    def add_latest_update(latest_update, doc, incNo, incPrio, chFile):
            if len(latest_update) != 3:
                raise ValueError
            else:
                previous_update = table.cell(12,1).text
                table.cell(13,1).text = latest_update[0] + ' - ' + latest_update[1] + '\n\n' + previous_update
                save_file(doc, incNo, incPrio, "UPDATE", chFile)
            

    chFile = choose_file()
    
    latest_update = pyperclip.paste()
    latest_update = latest_update.split('/nextEl,')
    
    
    if chFile == "output.docx":
        doc = docx.Document('output.docx')
        table = doc.tables[0]
    else:
        vbs_in(os.path.join('template', 'vbs', 'in.vbs ' + chFile))
        doc = docx.Document(os.path.join('template', 'vbs', 'temp.docx'))
        table = doc.tables[0]
        
# If there's another table in *msg file print error
    try:
        incNo = table.cell(2,1).text
        incPrio = table.cell(4,1).text[1]
        
        if x == '11':
            val = 'FF0000'
            filling(val)
            save_file(doc, incPrio, incNo, "INITIAL", chFile)
        elif x == '12':
            val = 'FFC000'
            table.cell(1,0).text = table.cell(1,0).text.replace('Initial', 'Update')
            table.cell(14,1).text = "Update"
            filling(val)
            save_file(doc, incPrio, incNo, "UPDATE", chFile)
        elif x == '13':
            val = '00B050'
            table.cell(1,0).text = table.cell(1,0).text.replace('Initial', 'Final')
            table.cell(4,3).text = "Resolved"
            table.cell(14,1).text = "Resolved"
            filling(val)
            save_file(doc, incPrio, incNo, "FINAL", chFile)
        elif x == '21':
            val = 'FF0000'
            filling(val)
            add_latest_update(latest_update, doc, incNo, incPrio, chFile)
        elif x == '22':
            val = 'FFC000'
            table.cell(1,0).text = table.cell(1,0).text.replace('Initial', 'Update')
            filling(val)
            add_latest_update(latest_update, doc, incNo, incPrio, chFile)
        elif x == '23':
            val = '00B050'
            table.cell(1,0).text = table.cell(1,0).text.replace('Initial', 'Final')
            table.cell(4,3).text = "Resolved"
            table.cell(14,1).text = "Resolved"
            filling(val)
            add_latest_update(latest_update, doc, incNo, incPrio, chFile)
        elif x == '24':
            add_latest_update(latest_update, doc, incNo, incPrio, chFile)
        else:
            print('\n!!!---You must input a number between 1 and 5---!!!\n')
            
    except IndexError:
        print('\n!!!---Make sure that in the file you are choosing is ONLY notification table---!!!\n')
    except ValueError:
        print('\n!!!---Invalid data format, press ALT+6 in SNow and try again---!!!\n')
        
    os.remove(os.path.join(os.getcwd(), "template", "vbs", "temp.docx"))
        
